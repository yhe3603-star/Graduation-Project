import os
from docx import Document

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                
    return '\n'.join(full_text)

folder_path = r'C:\Users\Administrator\Desktop\Graduation Project\2200470227谢世隆'
files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]

for file in files:
    try:
        path = os.path.join(folder_path, file)
        text = extract_text_from_docx(path)
        out_path = os.path.join(folder_path, file.replace('.docx', '.txt'))
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"Extracted {file}")
    except Exception as e:
        print(f"Failed to extract {file}: {e}")
